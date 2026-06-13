#!/usr/bin/env python3
"""Generate case_lookup_results.docx for the UK case lookup task."""
import docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/Users/a1234/.verdent/verdent-projects/case-annotation-review/case_lookup_results.docx"


def add_hyperlink(paragraph, url, text):
    """Insert a clickable hyperlink into a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def field_paragraph(doc, label, value, url=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"  {label:<10}")
    run.bold = True
    run.font.size = Pt(10)
    if url:
        add_hyperlink(p, url, value)
    else:
        r = p.add_run(value)
        r.font.size = Pt(10)
    return p


# Each case: name, citation, court, judge, date, parties, link, summary
# link=None means no public link found (note placed in summary/link field)
CASES = [
    {
        "n": 1,
        "name": "CDW Ltd v Bird & Anor",
        "citation": "[2021] EWHC 3665 (QB)",
        "court": "High Court of Justice, Queen's Bench Division (England & Wales)",
        "judge": "Heard before a QB judge (committal/freezing injunction proceedings)",
        "date": "2021 (judgment handed down December 2021)",
        "parties": "CDW Ltd (Claimant) v Bird & Anor (Defendants)",
        "link": "https://www.bailii.org/ew/cases/EWHC/QB/2021/3665.html",
        "link_note": "BAILII (stable). Not published on Find Case Law.",
        "summary": "Commercial fraud / freezing injunction proceedings in which the claimant company pursued a former employee and an associated party for misappropriated funds. The neutral citation [2021] EWHC 3665 (QB) is referenced by counsel (Blackstone Chambers); the full transcript is also available via CaseMine (subscription).",
    },
    {
        "n": 2,
        "name": "Davies v Carter",
        "citation": "[2021] EWHC 3021 (QB)",
        "court": "High Court of Justice, Queen's Bench Division (Media & Communications List)",
        "judge": "Trial judge sitting in the Media and Communications List",
        "date": "November 2021",
        "parties": "Terri Davies (Claimant) v Gavin Carter (Defendant)",
        "link": "https://www.bailii.org/ew/cases/EWHC/QB/2021/3021.html",
        "link_note": "BAILII (stable). Case analysis PDF also at Carter-Ruck.",
        "summary": "Libel and harassment claim arising from a near three-year social media campaign by the defendant, who irrationally convinced himself of the claimant's involvement in an allegedly defective service. The court granted an injunction and awarded GBP 35,000 in damages.",
    },
    {
        "n": 3,
        "name": "Day v Chivers",
        "citation": "No public neutral citation located",
        "court": "England and Wales High Court (Queen's Bench Division) - per CaseMine classification",
        "judge": "Not identified from public sources",
        "date": "December 2020 (preliminary issue ruling)",
        "parties": "Day (Plaintiff) v Chivers (Defendant)",
        "link": "https://www.casemine.com/judgement/uk/5fe970962c94e011bbd49f35",
        "link_note": "CaseMine (subscription) only; no BAILII / Find Case Law page found.",
        "summary": "Libel action over several Facebook posts published by the defendant as a business relationship deteriorated. The court determined, as a preliminary issue, the single natural and ordinary meaning of the posts and held each to be defamatory at common law. No public neutral citation was traced; appears unreported on free databases.",
    },
    {
        "n": 4,
        "name": "Drummond v Keolis Amey Docklands Ltd",
        "citation": "[2023] EWHC 853 (KB)",
        "court": "High Court of Justice, King's Bench Division (England & Wales)",
        "judge": "King's Bench Division judge",
        "date": "2023",
        "parties": "Drummond (Claimant) v Keolis Amey Docklands Ltd (Defendant)",
        "link": "https://caselaw.nationalarchives.gov.uk/ewhc/kb/2023/853",
        "link_note": "Find Case Law (The National Archives) - official, stable.",
        "summary": "Personal injury claim against the operator of the Docklands Light Railway. Reported on Find Case Law with the official neutral citation [2023] EWHC 853 (KB).",
    },
    {
        "n": 5,
        "name": "Grimshaw v Hudson",
        "citation": "No public neutral citation located",
        "court": "England and Wales High Court (Queen's Bench Division) - per CaseMine classification",
        "judge": "Not identified from public sources",
        "date": "Heard February 2021 (settlement approval)",
        "parties": "Grimshaw (Claimant, by litigation friend) v Hudson (Defendant)",
        "link": "https://www.casemine.com/judgement/uk/61a0cf0eb50db90e285bb7d7",
        "link_note": "CaseMine (subscription) only; no BAILII / Find Case Law page found.",
        "summary": "Approval of a personal-injury / clinical settlement of GBP 700,000 without admission of liability, under the court's inherent jurisdiction; causation and quantum had been contested. No public neutral citation traced; appears unreported on free databases.",
    },
    {
        "n": 6,
        "name": "JXM v An NHS Trust",
        "citation": "[2020] EWHC 919 (QB)",
        "court": "High Court of Justice, Queen's Bench Division (England & Wales)",
        "judge": "Mr Justice Martin Spencer",
        "date": "2020",
        "parties": "JXM (Claimant, anonymised, by litigation friend) v An NHS Trust (Defendant, anonymised)",
        "link": "https://caselaw.nationalarchives.gov.uk/ewhc/qb/2020/919",
        "link_note": "Find Case Law (The National Archives) - official, stable. Parties anonymised by court order.",
        "summary": "Clinical negligence claim where the defendant Trust admitted liability for failing to detect non-accidental rib fractures in an infant, who later suffered devastating brain damage. The judgment concerns approval of a GBP 600,000 interim payment and protective trust arrangements. Parties anonymised to protect the child.",
    },
    {
        "n": 7,
        "name": "Loughran (Paul) v Piney Rentals Limited & F5 Property Limited",
        "citation": "[2017] NICty 2",
        "court": "County Court for Northern Ireland (sitting in Belfast)",
        "judge": "District Judge Gilpin",
        "date": "21 December 2017",
        "parties": "Paul Loughran (Plaintiff) v Piney Rentals Limited (1st Defendant) & F5 Property Limited (2nd Defendant)",
        "link": "https://www.judiciaryni.uk/judicial-decisions/2017-nicty-2",
        "link_note": "Judiciary NI (official) - PDF also available there.",
        "summary": "A student tenant recovered a GBP 30 administration fee charged by letting agents, the court holding the fee void as commission under the Commission on Disposal of Lands (Northern Ireland) Order 1986. F5 Property Limited was ordered to repay.",
    },
    {
        "n": 8,
        "name": "Mahmood v Liverpool Victoria Insurance Company Ltd",
        "citation": "[2023] EW Misc 6 (CC)",
        "court": "County Court (England & Wales) - Circuit Judge",
        "judge": "Circuit Judge (county court)",
        "date": "6 July 2023",
        "parties": "Mahmood (Claimant) v Liverpool Victoria Insurance Company Ltd (Defendant)",
        "link": "https://www.casemine.com/judgement/uk/64a709436657545a79eeedd0",
        "link_note": "CaseMine (full judgment). Also reported on Find Case Law under EW Misc 6 (CC); not freely accessible via BAILII.",
        "summary": "Taxi credit-hire dispute following an RTC on 22 November 2020. The court restricted recoverable damages, finding a reasonable hire period of 22 days rather than the 124 days claimed, and limited credit-hire to loss-of-profit measure on the evidence.",
    },
    {
        "n": 9,
        "name": "Mirza v Farooqui & Anor",
        "citation": "[2021] EWHC 532 (QB)",
        "court": "High Court of Justice, Queen's Bench Division (Media & Communications List)",
        "judge": "Mrs Justice Collins Rice",
        "date": "9 March 2021",
        "parties": "Mirza (Claimant) v Farooqui & Anor (Defendants)",
        "link": "https://www.bailii.org/ew/cases/EWHC/QB/2021/532.html",
        "link_note": "BAILII (stable). Case No QB-2019-003019.",
        "summary": "Libel claim over a defamatory online newspaper article. Following default judgment, the court assessed damages, awarding the claimant GBP 75,000 plus ancillary relief for reputational and emotional harm.",
    },
    {
        "n": 10,
        "name": "NCL v MME",
        "citation": "[2020] EWHC 2679 (QB)",
        "court": "High Court of Justice, Queen's Bench Division (England & Wales)",
        "judge": "Mr Justice Lavender",
        "date": "2020",
        "parties": "NCL (Claimant, anonymised) v MME (Defendant, anonymised)",
        "link": "https://caselaw.nationalarchives.gov.uk/ewhc/qb/2020/2679",
        "link_note": "Find Case Law (The National Archives) - official, stable. Parties anonymised by initials.",
        "summary": "Anonymised proceedings before Mr Justice Lavender. Reported on Find Case Law with neutral citation [2020] EWHC 2679 (QB); the parties are identified only by initials, indicating anonymity protection.",
    },
    {
        "n": 11,
        "name": "Pass v Ministry of Defence",
        "citation": "[2021] EWHC 243 (QB)",
        "court": "High Court of Justice, Queen's Bench Division (England & Wales)",
        "judge": "Mr Justice Fordham",
        "date": "2021 (pre-trial review; trial fixed 1 March 2021)",
        "parties": "Pass (Claimant) v Ministry of Defence (Defendant)",
        "link": "https://caselaw.nationalarchives.gov.uk/ewhc/qb/2021/243",
        "link_note": "Find Case Law (The National Archives) - official, stable.",
        "summary": "Personal injury claim by a serving soldier alleging negligent delay in diagnosing and treating a spinal tumour, causing permanent disability and loss of his military career. Breach of duty was admitted; this judgment is a pre-trial review on evidential 'new inputs' with causation and quantum contested.",
    },
    {
        "n": 12,
        "name": "Robinson v Barker & Anor",
        "citation": "No public neutral citation located",
        "court": "England and Wales High Court (Queen's Bench Division) - per CaseMine classification",
        "judge": "Not identified from public sources",
        "date": "18 November 2020",
        "parties": "Robinson (Claimant) v Barker & Anor (Defendants)",
        "link": "https://www.casemine.com/judgement/uk/5fbb487e2c94e008949a1346",
        "link_note": "CaseMine (subscription) only; no BAILII / Find Case Law page found.",
        "summary": "Liability dispute arising from a road traffic collision on 15 December 2017 at Lancashire Hill, Stockport. No public neutral citation traced; appears unreported on free databases (CaseMine record only).",
    },
    {
        "n": 13,
        "name": "Sivananthan v Vasikaran",
        "citation": "[2022] EWHC 2938 (KB)",
        "court": "High Court of Justice, King's Bench Division (Media & Communications List)",
        "judge": "Mrs Justice Collins Rice",
        "date": "2022",
        "parties": "Sivananthan (Claimant) v Vasikaran (Defendant)",
        "link": "https://www.bailii.org/ew/cases/EWHC/KB/2022/2938.html",
        "link_note": "BAILII (stable). An earlier related ruling is [2022] EWHC 837 (QB).",
        "summary": "Defamation claim concerning WhatsApp / social-media messages, addressing serious harm to reputation under the Defamation Act 2013. Neutral citation [2022] EWHC 2938 (KB).",
    },
    {
        "n": 14,
        "name": "Spaul & Anor v Southfields Solicitors Ltd",
        "citation": "[2020] EWHC 1166 (QB)",
        "court": "High Court of Justice, Queen's Bench Division (England & Wales)",
        "judge": "Master Sullivan",
        "date": "14 May 2020",
        "parties": "Ashfaq Ahmed Spaul & Anor (Claimants) v Southfields Solicitors Ltd (Defendant)",
        "link": "https://www.bailii.org/ew/cases/EWHC/QB/2020/1166.html",
        "link_note": "BAILII (stable). Coverage also in Law Society Gazette.",
        "summary": "Claim for delivery up of conveyancing/refinancing files against the claimant's former solicitors. Following trial, the court dismissed the claim, rejecting the allegation that the firm had fabricated documents, and accepted the files had been loaned back to the client.",
    },
    {
        "n": 15,
        "name": "Walker v. Smith",
        "citation": "NOT FOUND",
        "court": "-",
        "judge": "-",
        "date": "-",
        "parties": "-",
        "link": None,
        "link_note": "No UK judgment identified.",
        "summary": "No UK reported judgment could be identified. The name is extremely common; searches returned only US (US Supreme Court 1804; 5th Circuit 2020), Californian, Jamaican and Australian (Victoria [2023] VSCA 61) cases. More identifying information (citation, year, court, or subject matter) is needed to locate a specific UK decision. See 'Not Found' list.",
    },
    {
        "n": 16,
        "name": "Wilson & Ors v Bayer Pharma AG & Ors",
        "citation": "[2023] EWHC 1282 (KB)",
        "court": "High Court of Justice, King's Bench Division (England & Wales)",
        "judge": "Mrs Justice Yip",
        "date": "26 May 2023",
        "parties": "Wilson & Others (Claimants) v Bayer Pharma AG & Others (Defendants)",
        "link": "https://www.bailii.org/ew/cases/EWHC/KB/2023/1282.html",
        "link_note": "BAILII (stable). Group litigation (Primodos hormone pregnancy test).",
        "summary": "Group litigation by claimants alleging injuries from Bayer's hormone pregnancy test (Primodos). The judgment of Mrs Justice Yip addresses limitation and case-management/strike-out issues in the group action. Neutral citation [2023] EWHC 1282 (KB).",
    },
]

NOT_FOUND = [
    {
        "name": "Walker v. Smith",
        "tried": "Walker v Smith judgment EWHC bailii; \"Walker v Smith\" EWHC OR county court UK England judgment personal injury OR credit hire OR defamation; \"Walker v Smith\" EWHC casemine England and Wales judgment United Kingdom",
        "reason": "Name too generic. Only non-UK results found (US Supreme Court 4 U.S. 389 (1804); US 5th Cir. 2020; California Court of Appeal 1958; Jamaica vLex; Australia Smith v Walker [2023] VSCA 61). No matching England & Wales / NI / Scotland judgment located. Needs neutral citation, year, court or subject matter to disambiguate.",
    },
]


def main():
    doc = Document()

    # base style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Title
    h = doc.add_heading("UK Case Law Lookup Results", level=0)
    sub = doc.add_paragraph()
    r = sub.add_run("Original judgments located via BAILII, Find Case Law (The National Archives), "
                    "Judiciary NI and CaseMine. 17 cases requested.")
    r.italic = True
    r.font.size = Pt(9)

    found_count = sum(1 for c in CASES if c["citation"] != "NOT FOUND")
    stat = doc.add_paragraph()
    rs = stat.add_run(f"Found: {found_count}    |    Not found: {len(CASES) - found_count}    "
                      f"(Note: 4 cases - Day v Chivers, Grimshaw v Hudson, Robinson v Barker, and "
                      f"Mahmood - are reachable only via subscription databases, with no free "
                      f"BAILII/Find Case Law transcript; Day, Grimshaw and Robinson also have no "
                      f"traceable public neutral citation.)")
    rs.font.size = Pt(9)
    rs.bold = True

    doc.add_paragraph()

    for c in CASES:
        if c["citation"] == "NOT FOUND":
            continue
        head = doc.add_heading(level=1)
        run = head.add_run(f'Case {c["n"]}. {c["name"]}')
        run.font.size = Pt(13)

        field_paragraph(doc, "Citation:", c["citation"])
        field_paragraph(doc, "Court:", c["court"])
        field_paragraph(doc, "Judge:", c["judge"])
        field_paragraph(doc, "Date:", c["date"])
        field_paragraph(doc, "Parties:", c["parties"])
        if c["link"]:
            field_paragraph(doc, "Link:", c["link"], url=c["link"])
            note = doc.add_paragraph()
            note.paragraph_format.space_after = Pt(2)
            rn = note.add_run(f"             ({c['link_note']})")
            rn.italic = True
            rn.font.size = Pt(8)
            rn.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        else:
            field_paragraph(doc, "Link:", c["link_note"])
        field_paragraph(doc, "Summary:", c["summary"])
        doc.add_paragraph()

    # Not found section
    doc.add_page_break()
    nf_head = doc.add_heading("未找到清单 / Not Found", level=1)
    intro = doc.add_paragraph()
    ri = intro.add_run("The following case(s) could not be matched to a UK judgment with the "
                       "information provided.")
    ri.font.size = Pt(10)

    for c in NOT_FOUND:
        ph = doc.add_heading(level=2)
        ph.add_run(c["name"]).font.size = Pt(11)
        field_paragraph(doc, "Status:", "Not found / name too generic — more identifying information needed")
        field_paragraph(doc, "Searched:", c["tried"])
        field_paragraph(doc, "Reason:", c["reason"])
        doc.add_paragraph()

    doc.save(OUT)
    print("Saved:", OUT)


if __name__ == "__main__":
    main()
